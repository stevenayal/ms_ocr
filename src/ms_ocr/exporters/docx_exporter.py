"""Export to DOCX format."""

from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor

from ms_ocr.layout.rules import LayoutElement
from ms_ocr.tables.tables_extractor import Table


class DocxExporter:
    """Export document to DOCX."""

    def __init__(self):
        """Initialize DOCX exporter."""
        pass

    def export(
        self,
        elements: List[LayoutElement],
        tables: List[Table],
        output_path: Path,
        title: Optional[str] = None,
    ):
        """Export to DOCX file.

        Args:
            elements: Layout elements
            tables: Extracted tables
            output_path: Output file path
            title: Document title
        """
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Add title
        if title:
            doc.add_heading(title, level=0)

        # Add elements
        for element in elements:
            self._add_element(doc, element)

        # Add tables
        if tables:
            doc.add_page_break()
            doc.add_heading("Tables", level=1)

            for i, table in enumerate(tables):
                doc.add_heading(f"Table {i + 1} (Page {table.page_num + 1})", level=2)
                self._add_table(doc, table)

        # Save
        doc.save(output_path)

    def _add_element(self, doc: Document, element: LayoutElement):
        """Add layout element to document.

        Args:
            doc: Document object
            element: Layout element
        """
        if element.type == "title":
            doc.add_heading(element.text, level=0)

        elif element.type == "heading":
            level = min(element.level or 1, 9)
            doc.add_heading(element.text, level=level)

        elif element.type == "paragraph":
            doc.add_paragraph(element.text)

        elif element.type == "list_item":
            # Remove bullet/number prefix
            text = element.text.lstrip("•●○■□▪▫-*0123456789. \t")
            doc.add_paragraph(text, style="List Bullet")

        else:
            doc.add_paragraph(element.text)

    def _add_table(self, doc: Document, table: Table):
        """Add table to document.

        Args:
            doc: Document object
            table: Table object
        """
        if table.data.empty:
            return

        # Create table
        df = table.data
        rows, cols = df.shape

        doc_table = doc.add_table(rows=rows + 1, cols=cols)
        doc_table.style = "Light Grid Accent 1"

        # Add header
        for i, col in enumerate(df.columns):
            cell = doc_table.rows[0].cells[i]
            cell.text = str(col)

        # Add data
        for i, row in df.iterrows():
            for j, value in enumerate(row):
                cell = doc_table.rows[i + 1].cells[j]
                cell.text = str(value)

        doc.add_paragraph()  # Add space after table
